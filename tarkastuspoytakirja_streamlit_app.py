import base64
import json
from datetime import date, datetime
from io import BytesIO
from typing import Dict, List

import streamlit as st
from PIL import Image
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


st.set_page_config(
    page_title="Tarkastuspöytäkirja",
    page_icon="📄",
    layout="wide",
)

LOGO_PATH = "mirus_logo.png"
COPYRIGHT_TEXT = "Tekijänoikeudet kuuluvat Mirus Electrum Oy:lle."
WEBSITE_URL = "https://mirus-electrum.fi"

STATUS_OPTIONS = [
    "OK",
    "Huomio",
    "Korjattava",
    "Ei tarkastettu",
    "Ei sovellu",
]

RISK_OPTIONS = ["Matala", "Keskitaso", "Korkea", "Välitön vaara"]

SECTION_DEFINITIONS = {
    "Yleistarkastus": [
        "Kohteeseen pääsy ja kulkureitit turvalliset",
        "Tarkastettava tila siisti ja esteetön",
        "Lukitus, avainhallinta ja pääsyrajoitukset kunnossa",
        "Varoitus- ja ohjekilvet näkyvissä",
        "Merkinnät ajantasalla",
        "Laitteiden käyttöohjeet saatavilla",
        "Valaistus riittävä ja toimiva",
        "Poikkeavat hajut, äänet tai lämpenemät",
        "Kosteus-, pöly- tai vesivuotojälkiä havaittavissa",
        "Tilan palokatkot",
        "Huolto- ja kunnossapitosuunnitelman toteutuminen",
        "Suunnitelmat vastaavat toteutusta",
    ],
    "Keskijännitekojeisto": [
        "Kennot, ovet ja lukitukset ehjät",
        "Kojeiston merkinnät ja tunnukset selkeät",
        "Erotin-/katkaisijatilojen asennot ja osoittimet kunnossa",
        "Mekaaniset lukitukset toimivat",
        "Suojareleiden ja mittalaitteiden tila normaali",
        "Hälytykset tai vikailmoitukset tarkastettu",
        "Käyttövipujen, varusteiden ja maadoitusvälineiden kunto",
        "Kiskosto- ja liitäntätiloissa ei näkyviä vaurioita",
        "Osittaispurkaus-, noki- tai valokaarijälkiä ei havaita",
        "Huolto- ja koestusmerkinnät ajan tasalla",
        "Suunnitelmat vastaavat toteutusta",
    ],
    "Muuntaja": [
        "Muuntajan yleiskunto ja puhtaus",
        "Merkinnät ja arvokilpi luettavissa",
        "Epänormaalia ääntä, tärinää tai hajua ei havaita",
        "Läpiviennit, eristimet ja liitokset ehjät",
        "Öljyvuotoja ei havaita",
        "Öljyn pinnankorkeus / ilmaisin normaali (jos soveltuu)",
        "Lämpötilan valvonta ja hälytykset kunnossa",
        "Jäähdytys, ilmanvaihto ja tilan lämpöolosuhteet kunnossa",
        "Maadoitukset ja potentiaalintasaukset kunnossa",
        "Suojaukset ja esteet asianmukaiset",
        "Suunnitelmat vastaavat toteutusta",
    ],
    "Pääkeskus": [
        "Keskuksen yleiskunto ja siisteys",
        "Keskuksen merkinnät ja ryhmätiedot ajan tasalla",
        "Ovet, suojaukset ja peitelevyt ehjät",
        "Näkyviä lämpenemä-, noki- tai vauriojälkiä ei havaita",
        "Johdotukset ja läpiviennit asianmukaiset",
        "Lukitus ja pääsy järjestetty",
        "Mittaukset ja indikoinnit normaalit",
        "Ylivirta- ja vikavirtasuojien merkinnät kunnossa",
        "Varasyötöt / varavoima merkinnöin ja ohjein tunnistettavissa",
        "Pää- ja käyttökatkaisijoiden käyttömahdollisuus esteetön",
        "Suunnitelmat vastaavat toteutusta",
    ],
    "ATEX-tila": [
        "ATEX-luokitus ja tilaluokitusmerkinnät näkyvissä",
        "Laitteiden Ex-merkinnät asianmukaiset",
        "Kotelointiluokat ja mekaaninen kunto kunnossa",
        "Läpiviennit ja tiivisteet ehjät",
        "Potentiaalintasaus ja maadoitukset kunnossa",
        "Pölyn / kaasun kertymää ei havaita poikkeavasti",
        "Huolto- ja tarkastusmerkinnät ajan tasalla",
        "Muutostöiden dokumentointi saatavilla",
        "Sytytyslähteiden hallinta toteutuu",
        "Tilassa työskentelyn erityisohjeet saatavilla",
        "Huolto- ja kunnossapitosuunnitelman toteutuminen",
        "Toiminta vastaa RSA-asiakirjaa",
        "Suunnitelmat vastaavat toteutusta",
    ],
    "Muut tilat / huollon havainnot": [
        "Muut havainnot kiinteistöstä/laitteistosta",
        "Tehdyistä muutoksista on olemassa käyttöönottotarkastuspöytäkirja",
        "Kiinteistön muut järjestelmät, jotka vaativat tarkastusta",
        "Suunnitelmat vastaavat toteutusta",
    ],
}

DEFAULT_ENABLED = ["Yleistarkastus", "Keskijännitekojeisto", "Muuntaja 1", "Pääkeskus"]


def init_state() -> None:
    if "selected_sections" not in st.session_state:
        st.session_state.selected_sections = DEFAULT_ENABLED.copy()
    if "transformer_count" not in st.session_state:
        st.session_state.transformer_count = 1
    if "pending_loaded_json" not in st.session_state:
        st.session_state.pending_loaded_json = None
    if "json_loaded_message" not in st.session_state:
        st.session_state.json_loaded_message = ""


def load_logo_bytes() -> bytes:
    try:
        with open(LOGO_PATH, "rb") as f:
            return f.read()
    except FileNotFoundError:
        return b""


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def get_transformer_sections() -> List[str]:
    count = max(0, int(st.session_state.get("transformer_count", 1)))
    return [f"Muuntaja {idx}" for idx in range(1, count + 1)]


def get_available_sections() -> List[str]:
    return [
        "Yleistarkastus",
        "Keskijännitekojeisto",
        *get_transformer_sections(),
        "Pääkeskus",
        "ATEX-tila",
        "Muut tilat / huollon havainnot",
    ]


def sync_selected_sections() -> None:
    available = set(get_available_sections())
    selected = st.session_state.get("selected_sections", [])
    filtered = [sec for sec in selected if sec in available]
    if not filtered:
        filtered = [sec for sec in DEFAULT_ENABLED if sec in available]
    st.session_state.selected_sections = filtered


def sync_sections_from_tarkastustyyppi() -> None:
    tarkastuksen_tyyppi = st.session_state.get("tarkastuksen_tyyppi", "Muuntamotarkastus")
    current = set(st.session_state.get("selected_sections", []))
    if tarkastuksen_tyyppi in ["ATEX", "Muuntamotarkastus +ATEX"]:
        current.add("ATEX-tila")
    st.session_state.selected_sections = [sec for sec in get_available_sections() if sec in current]


def get_base_section_name(section_name: str) -> str:
    if section_name.startswith("Muuntaja "):
        return "Muuntaja"
    return section_name


def make_key(section: str, item_idx: int, suffix: str) -> str:
    return f"{section}_{item_idx}_{suffix}".replace(" ", "_").replace("/", "_")


def resize_image_bytes(file_bytes: bytes, max_width: int = 1600, jpeg_quality: int = 82) -> tuple[bytes, str]:
    image = Image.open(BytesIO(file_bytes))
    if image.mode not in ("RGB", "L"):
        image = image.convert("RGB")

    width, height = image.size
    if width > max_width:
        new_height = int(height * (max_width / width))
        image = image.resize((max_width, new_height))

    output = BytesIO()
    image.save(output, format="JPEG", quality=jpeg_quality, optimize=True)
    return output.getvalue(), "image/jpeg"


def encode_uploaded_files(uploaded_files) -> List[Dict]:
    images = []
    if not uploaded_files:
        return images

    for uploaded_file in uploaded_files:
        file_bytes = uploaded_file.getvalue()
        resized_bytes, mime_type = resize_image_bytes(file_bytes)
        images.append(
            {
                "name": uploaded_file.name,
                "mime_type": mime_type,
                "data_base64": base64.b64encode(resized_bytes).decode("utf-8"),
            }
        )
    return images


def decode_image_bytes(image_dict: Dict) -> bytes:
    if not image_dict or not image_dict.get("data_base64"):
        return b""
    return base64.b64decode(image_dict["data_base64"])


def image_to_data_uri(image_dict: Dict) -> str:
    if not image_dict:
        return ""
    mime_type = image_dict.get("mime_type", "image/jpeg")
    return f"data:{mime_type};base64,{image_dict.get('data_base64', '')}"


def render_check_section(section_name: str, items: List[str]) -> None:
    st.subheader(section_name)
    st.caption("Merkitse havainto jokaisesta kohdasta ja lisää tarvittaessa kommentti ja yksi tai useampi valokuva.")

    for idx, item in enumerate(items):
        with st.container(border=True):
            col1, col2 = st.columns([2.4, 1])
            with col1:
                st.markdown(f"**{idx + 1}. {item}**")
                st.text_area(
                    "Kommentti",
                    key=make_key(section_name, idx, "comment"),
                    height=80,
                    placeholder="Kirjaa havainto, mittaustulos, poikkeama tai lisätieto...",
                    label_visibility="collapsed",
                )
                uploaded_images = st.file_uploader(
                    "Liitä havaintokuvia",
                    type=["jpg", "jpeg", "png", "webp"],
                    key=make_key(section_name, idx, "image_uploader"),
                    accept_multiple_files=True,
                    help="Voit liittää useita kuvia tälle tarkastuskohdalle. Kuvat pienennetään automaattisesti.",
                )
                if uploaded_images:
                    image_list = encode_uploaded_files(uploaded_images)
                    st.session_state[make_key(section_name, idx, "image_data")] = image_list

                saved_images = st.session_state.get(make_key(section_name, idx, "image_data"), [])
                if saved_images:
                    st.markdown("**Tallennetut kuvat**")
                    preview_cols = st.columns(3)
                    for image_idx, saved_image in enumerate(saved_images):
                        with preview_cols[image_idx % 3]:
                            st.image(
                                decode_image_bytes(saved_image),
                                caption=saved_image.get("name", f"Kuva {image_idx + 1}"),
                                width=220,
                            )
                    if st.button("Poista kaikki kuvat", key=make_key(section_name, idx, "remove_image")):
                        st.session_state[make_key(section_name, idx, "image_data")] = []
                        st.rerun()

            with col2:
                st.selectbox(
                    "Tila",
                    STATUS_OPTIONS,
                    key=make_key(section_name, idx, "status"),
                )
                st.selectbox(
                    "Riskitaso",
                    RISK_OPTIONS,
                    key=make_key(section_name, idx, "risk"),
                )
                st.text_input(
                    "Toimenpide / vastuuhenkilö",
                    key=make_key(section_name, idx, "action"),
                    placeholder="Esim. huolto / kiinteistöpäällikkö",
                )


def collect_section_data(section_name: str, items: List[str]) -> List[Dict]:
    data = []
    for idx, item in enumerate(items):
        data.append(
            {
                "kohta": item,
                "tila": st.session_state.get(make_key(section_name, idx, "status"), "Ei tarkastettu"),
                "riskitaso": st.session_state.get(make_key(section_name, idx, "risk"), "Matala"),
                "kommentti": st.session_state.get(make_key(section_name, idx, "comment"), ""),
                "toimenpide": st.session_state.get(make_key(section_name, idx, "action"), ""),
                "kuvat": st.session_state.get(make_key(section_name, idx, "image_data"), []),
            }
        )
    return data


def section_summary_rows(section_name: str, rows: List[Dict]) -> str:
    html_rows = ""
    for row in rows:
        status_class = (
            "status-ok" if row["tila"] == "OK" else "status-warn" if row["tila"] in ["Huomio", "Ei tarkastettu"] else "status-bad"
        )
        image_html = "-"
        if row.get("kuvat"):
            image_tags = []
            for image_dict in row["kuvat"]:
                image_tags.append(
                    f'<img src="{image_to_data_uri(image_dict)}" style="max-width:220px; max-height:180px; margin:4px; border:1px solid #cbd5e1; border-radius:6px;" />'
                )
            image_html = "".join(image_tags)

        html_rows += f"""
        <tr>
            <td>{row['kohta']}</td>
            <td><span class="pill {status_class}">{row['tila']}</span></td>
            <td>{row['riskitaso']}</td>
            <td>{row['kommentti'] or '-'}</td>
            <td>{row['toimenpide'] or '-'}</td>
            <td>{image_html}</td>
        </tr>
        """

    return f"""
    <h2>{section_name}</h2>
    <table>
        <thead>
            <tr>
                <th>Kohta</th>
                <th>Tila</th>
                <th>Riskitaso</th>
                <th>Kommentti</th>
                <th>Toimenpide / vastuuhenkilö</th>
                <th>Valokuvat</th>
            </tr>
        </thead>
        <tbody>
            {html_rows}
        </tbody>
    </table>
    """


def build_report(data: Dict) -> str:
    section_html = ""
    for section_name, rows in data["tarkastusosiot"].items():
        section_html += section_summary_rows(section_name, rows)

    logo_bytes = load_logo_bytes()
    logo_html = ""
    if logo_bytes:
        encoded_logo = base64.b64encode(logo_bytes).decode("utf-8")
        logo_html = f'<img src="data:image/png;base64,{encoded_logo}" style="max-height:90px; margin-bottom:18px;" />'

    return f"""
    <!DOCTYPE html>
    <html lang="fi">
    <head>
        <meta charset="utf-8" />
        <title>Tarkastuspöytäkirja - {data['kohde']['kohteen_nimi']}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 30px; color: #1f2937; }}
            h1, h2, h3 {{ color: #111827; }}
            .meta {{ display: grid; grid-template-columns: 1fr 1fr; gap: 10px 30px; margin-bottom: 24px; }}
            .box {{ background: #f8fafc; border: 1px solid #dbe2ea; border-radius: 8px; padding: 14px; margin-bottom: 18px; }}
            table {{ width: 100%; border-collapse: collapse; margin-bottom: 24px; }}
            th, td {{ border: 1px solid #d1d5db; padding: 8px; vertical-align: top; text-align: left; font-size: 13px; }}
            th {{ background: #f3f4f6; }}
            .pill {{ display: inline-block; padding: 4px 8px; border-radius: 999px; font-size: 12px; font-weight: bold; }}
            .status-ok {{ background: #dcfce7; color: #166534; }}
            .status-warn {{ background: #fef3c7; color: #92400e; }}
            .status-bad {{ background: #fee2e2; color: #991b1b; }}
            .footer {{ margin-top: 30px; font-size: 12px; color: #6b7280; }}
            @media print {{ body {{ margin: 12mm; }} .no-print {{ display: none; }} }}
        </style>
    </head>
    <body>
        {logo_html}
        <h1>Tarkastuspöytäkirja</h1>

        <div class="box">
            <h3>Kohdetiedot</h3>
            <div class="meta">
                <div><strong>Kohde:</strong> {data['kohde']['kohteen_nimi']}</div>
                <div><strong>Osoite:</strong> {data['kohde']['osoite']}</div>
                <div><strong>Tarkastuspäivä:</strong> {data['kohde']['tarkastuspaiva']}</div>
                <div><strong>Tarkastaja:</strong> {data['kohde']['tarkastaja']}</div>
                <div><strong>Käytönjohtaja:</strong> {data['kohde']['kaytonjohtaja']}</div>
                <div><strong>Kiinteistön yhteyshenkilö:</strong> {data['kohde']['yhteyshenkilo']}</div>
                <div><strong>Suurin nimellisjännite:</strong> {data['kohde']['suurin_nimellisjannite']}</div>
                <div><strong>Kohteen tyyppi:</strong> {data['kohde']['kohteen_tyyppi']}</div>
                <div><strong>Tarkastuksen tyyppi:</strong> {data['kohde']['tarkastuksen_tyyppi']}</div>
                <div><strong>Sää / olosuhteet:</strong> {data['kohde']['olosuhteet']}</div>
            </div>
        </div>

        <div class="box">
            <h3>Yhteenveto</h3>
            <p><strong>Kokonaisarvio:</strong> {data['yhteenveto']['kokonaisarvio']}</p>
            <p><strong>Välittömät toimenpiteet:</strong> {data['yhteenveto']['valittomat_toimet'] or '-'}</p>
            <p><strong>Muut toimenpiteet:</strong> {data['yhteenveto']['muut_toimet'] or '-'}</p>
            <p><strong>Seuraava tarkastus:</strong> {data['yhteenveto']['seuraava_tarkastus']}</p>
            <p><strong>Lisähuomiot:</strong> {data['yhteenveto']['lisahuomiot'] or '-'}</p>
        </div>

        {section_html}

        <div class="box">
            <h3>Kuittaukset</h3>
            <p><strong>Tarkastaja:</strong> {data['kuittaukset']['tarkastaja_kuittaus']}</p>
            <p><strong>Vastaanottaja:</strong> {data['kuittaukset']['vastaanottaja_kuittaus']}</p>
        </div>

        <div class="footer">
            <div>Luotu {datetime.now().strftime('%d.%m.%Y %H:%M')} Streamlit-tarkastussovelluksella.</div>
            <div style="margin-top:8px;">{COPYRIGHT_TEXT}</div>
            <div style="margin-top:4px;"><a href="{WEBSITE_URL}">{WEBSITE_URL}</a></div>
        </div>
    </body>
    </html>
    """


def build_pdf(report_data: Dict) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=24,
        leftMargin=24,
        topMargin=24,
        bottomMargin=24,
    )
    styles = getSampleStyleSheet()
    story = []

    logo_bytes = load_logo_bytes()
    if logo_bytes:
        try:
            story.append(RLImage(BytesIO(logo_bytes), width=120, height=120))
            story.append(Spacer(1, 8))
        except Exception:
            pass

    story.append(Paragraph("Tarkastuspöytäkirja", styles["Title"]))
    story.append(Spacer(1, 12))

    kohde_rows = [
        ["Kohde", report_data["kohde"]["kohteen_nimi"]],
        ["Osoite", report_data["kohde"]["osoite"]],
        ["Tarkastuspäivä", report_data["kohde"]["tarkastuspaiva"]],
        ["Tarkastaja", report_data["kohde"]["tarkastaja"]],
        ["Käytönjohtaja", report_data["kohde"]["kaytonjohtaja"]],
        ["Kiinteistön yhteyshenkilö", report_data["kohde"]["yhteyshenkilo"]],
        ["Suurin nimellisjännite", report_data["kohde"]["suurin_nimellisjannite"]],
        ["Kohteen tyyppi", report_data["kohde"]["kohteen_tyyppi"]],
        ["Tarkastuksen tyyppi", report_data["kohde"]["tarkastuksen_tyyppi"]],
        ["Sää / olosuhteet", report_data["kohde"]["olosuhteet"]],
    ]
    kohde_table = Table(kohde_rows, colWidths=[150, 360])
    kohde_table.setStyle(
        TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ])
    )
    story.append(Paragraph("Kohdetiedot", styles["Heading2"]))
    story.append(kohde_table)
    story.append(Spacer(1, 12))

    story.append(Paragraph("Yhteenveto", styles["Heading2"]))
    for label, value in [
        ("Kokonaisarvio", report_data["yhteenveto"]["kokonaisarvio"]),
        ("Välittömät toimenpiteet", report_data["yhteenveto"]["valittomat_toimet"] or "-"),
        ("Muut toimenpiteet", report_data["yhteenveto"]["muut_toimet"] or "-"),
        ("Seuraava tarkastus", report_data["yhteenveto"]["seuraava_tarkastus"]),
        ("Lisähuomiot", report_data["yhteenveto"]["lisahuomiot"] or "-"),
    ]:
        story.append(Paragraph(f"<b>{label}:</b> {value}", styles["BodyText"]))
        story.append(Spacer(1, 6))

    for section_name, rows in report_data["tarkastusosiot"].items():
        story.append(Spacer(1, 8))
        story.append(Paragraph(section_name, styles["Heading2"]))
        table_data = [["Kohta", "Tila", "Riski", "Kommentti", "Toimenpide"]]
        for row in rows:
            table_data.append([
                row["kohta"],
                row["tila"],
                row["riskitaso"],
                row["kommentti"] or "-",
                row["toimenpide"] or "-",
            ])
        tbl = Table(table_data, repeatRows=1, colWidths=[150, 60, 60, 130, 110])
        tbl.setStyle(
            TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ])
        )
        story.append(tbl)

        for idx, row in enumerate(rows, start=1):
            if row.get("kuvat"):
                story.append(Spacer(1, 4))
                story.append(Paragraph(f"{section_name} / kohta {idx}: {row['kohta']} - valokuvat", styles["BodyText"]))
                for image_dict in row["kuvat"]:
                    img_bytes = decode_image_bytes(image_dict)
                    if img_bytes:
                        try:
                            story.append(RLImage(BytesIO(img_bytes), width=220, height=165))
                        except Exception:
                            story.append(Paragraph("Kuvaa ei voitu liittää PDF-raporttiin.", styles["BodyText"]))
                        story.append(Spacer(1, 4))
                story.append(Spacer(1, 6))

    story.append(Spacer(1, 12))
    story.append(Paragraph("Kuittaukset", styles["Heading2"]))
    story.append(Paragraph(f"<b>Tarkastaja:</b> {report_data['kuittaukset']['tarkastaja_kuittaus'] or '-'}", styles["BodyText"]))
    story.append(Paragraph(f"<b>Vastaanottaja:</b> {report_data['kuittaukset']['vastaanottaja_kuittaus'] or '-'}", styles["BodyText"]))
    story.append(Spacer(1, 10))
    story.append(Paragraph(COPYRIGHT_TEXT, styles["BodyText"]))
    story.append(Paragraph(f'<link href="{WEBSITE_URL}">{WEBSITE_URL}</link>', styles["BodyText"]))

    doc.build(story)
    return buffer.getvalue()


def build_word(report_data: Dict) -> bytes:
    doc = Document()

    logo_bytes = load_logo_bytes()
    if logo_bytes:
        try:
            doc.add_picture(BytesIO(logo_bytes), width=Inches(1.4))
        except Exception:
            pass

    doc.add_heading("Tarkastuspöytäkirja", level=0)

    doc.add_heading("Kohdetiedot", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in [
        ("Kohde", report_data["kohde"]["kohteen_nimi"]),
        ("Osoite", report_data["kohde"]["osoite"]),
        ("Tarkastuspäivä", report_data["kohde"]["tarkastuspaiva"]),
        ("Tarkastaja", report_data["kohde"]["tarkastaja"]),
        ("Käytönjohtaja", report_data["kohde"]["kaytonjohtaja"]),
        ("Kiinteistön yhteyshenkilö", report_data["kohde"]["yhteyshenkilo"]),
        ("Suurin nimellisjännite", report_data["kohde"]["suurin_nimellisjannite"]),
        ("Kohteen tyyppi", report_data["kohde"]["kohteen_tyyppi"]),
        ("Tarkastuksen tyyppi", report_data["kohde"]["tarkastuksen_tyyppi"]),
        ("Sää / olosuhteet", report_data["kohde"]["olosuhteet"]),
    ]:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    doc.add_heading("Yhteenveto", level=1)
    for label, value in [
        ("Kokonaisarvio", report_data["yhteenveto"]["kokonaisarvio"]),
        ("Välittömät toimenpiteet", report_data["yhteenveto"]["valittomat_toimet"] or "-"),
        ("Muut toimenpiteet", report_data["yhteenveto"]["muut_toimet"] or "-"),
        ("Seuraava tarkastus", report_data["yhteenveto"]["seuraava_tarkastus"]),
        ("Lisähuomiot", report_data["yhteenveto"]["lisahuomiot"] or "-"),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value))

    for section_name, rows in report_data["tarkastusosiot"].items():
        doc.add_heading(section_name, level=1)
        section_table = doc.add_table(rows=1, cols=5)
        section_table.style = "Table Grid"
        hdr = section_table.rows[0].cells
        hdr[0].text = "Kohta"
        hdr[1].text = "Tila"
        hdr[2].text = "Riski"
        hdr[3].text = "Kommentti"
        hdr[4].text = "Toimenpide"

        for rowdata in rows:
            row = section_table.add_row().cells
            row[0].text = rowdata["kohta"]
            row[1].text = rowdata["tila"]
            row[2].text = rowdata["riskitaso"]
            row[3].text = rowdata["kommentti"] or "-"
            row[4].text = rowdata["toimenpide"] or "-"

        for idx, rowdata in enumerate(rows, start=1):
            if rowdata.get("kuvat"):
                doc.add_paragraph(f"{section_name} / kohta {idx}: {rowdata['kohta']} - valokuvat")
                for image_dict in rowdata["kuvat"]:
                    img_bytes = decode_image_bytes(image_dict)
                    if img_bytes:
                        try:
                            doc.add_picture(BytesIO(img_bytes), width=Inches(3.2))
                        except Exception:
                            doc.add_paragraph("Kuvaa ei voitu liittää Word-raporttiin.")

    doc.add_heading("Kuittaukset", level=1)
    p1 = doc.add_paragraph()
    p1.add_run("Tarkastaja: ").bold = True
    p1.add_run(report_data["kuittaukset"]["tarkastaja_kuittaus"] or "-")

    p2 = doc.add_paragraph()
    p2.add_run("Vastaanottaja: ").bold = True
    p2.add_run(report_data["kuittaukset"]["vastaanottaja_kuittaus"] or "-")

    footer_para = doc.add_paragraph()
    footer_para.add_run(COPYRIGHT_TEXT)

    footer_link_para = doc.add_paragraph()
    add_hyperlink(footer_link_para, WEBSITE_URL, WEBSITE_URL)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def gather_form_data() -> Dict:
    selected_sections = st.session_state.get("selected_sections", [])
    tarkastuksen_tyyppi = st.session_state.get("tarkastuksen_tyyppi", "Muuntamotarkastus")

    if tarkastuksen_tyyppi in ["ATEX", "Muuntamotarkastus +ATEX"] and "ATEX-tila" not in selected_sections:
        selected_sections = selected_sections + ["ATEX-tila"]

    report_data = {
        "kohde": {
            "kohteen_nimi": st.session_state.get("kohteen_nimi", ""),
            "osoite": st.session_state.get("osoite", ""),
            "tarkastuspaiva": str(st.session_state.get("tarkastuspaiva", date.today())),
            "tarkastaja": st.session_state.get("tarkastaja", ""),
            "kaytonjohtaja": st.session_state.get("kaytonjohtaja", ""),
            "yhteyshenkilo": st.session_state.get("yhteyshenkilo", ""),
            "suurin_nimellisjannite": st.session_state.get("suurin_nimellisjannite", "20 kV"),
            "kohteen_tyyppi": st.session_state.get("kohteen_tyyppi", ""),
            "tarkastuksen_tyyppi": tarkastuksen_tyyppi,
            "olosuhteet": st.session_state.get("olosuhteet", ""),
        },
        "yhteenveto": {
            "kokonaisarvio": st.session_state.get("kokonaisarvio", ""),
            "valittomat_toimet": st.session_state.get("valittomat_toimet", ""),
            "muut_toimet": st.session_state.get("muut_toimet", ""),
            "seuraava_tarkastus": str(st.session_state.get("seuraava_tarkastus", date.today())),
            "lisahuomiot": st.session_state.get("lisahuomiot", ""),
        },
        "tarkastusosiot": {},
        "liitteet": [],
        "kuittaukset": {
            "tarkastaja_kuittaus": st.session_state.get("tarkastaja_kuittaus", ""),
            "vastaanottaja_kuittaus": st.session_state.get("vastaanottaja_kuittaus", ""),
        },
    }

    for section in selected_sections:
        base_section = get_base_section_name(section)
        report_data["tarkastusosiot"][section] = collect_section_data(section, SECTION_DEFINITIONS[base_section])

    return report_data


def load_json_to_state(data: Dict) -> None:
    kohde = data.get("kohde", {})
    yhteenveto = data.get("yhteenveto", {})
    kuittaukset = data.get("kuittaukset", {})

    st.session_state["kohteen_nimi"] = kohde.get("kohteen_nimi", "")
    st.session_state["osoite"] = kohde.get("osoite", "")
    st.session_state["tarkastaja"] = kohde.get("tarkastaja", "")
    st.session_state["kaytonjohtaja"] = kohde.get("kaytonjohtaja", "")
    st.session_state["yhteyshenkilo"] = kohde.get("yhteyshenkilo", "")
    st.session_state["suurin_nimellisjannite"] = kohde.get("suurin_nimellisjannite", "20 kV")
    st.session_state["kohteen_tyyppi"] = kohde.get("kohteen_tyyppi", "")
    st.session_state["tarkastuksen_tyyppi"] = kohde.get("tarkastuksen_tyyppi", "Muuntamotarkastus")
    st.session_state["olosuhteet"] = kohde.get("olosuhteet", "")

    try:
        st.session_state["tarkastuspaiva"] = datetime.strptime(
            kohde.get("tarkastuspaiva", str(date.today())),
            "%Y-%m-%d",
        ).date()
    except ValueError:
        st.session_state["tarkastuspaiva"] = date.today()

    try:
        st.session_state["seuraava_tarkastus"] = datetime.strptime(
            yhteenveto.get("seuraava_tarkastus", str(date.today())),
            "%Y-%m-%d",
        ).date()
    except ValueError:
        st.session_state["seuraava_tarkastus"] = date.today()

    st.session_state["kokonaisarvio"] = yhteenveto.get("kokonaisarvio", "")
    st.session_state["valittomat_toimet"] = yhteenveto.get("valittomat_toimet", "")
    st.session_state["muut_toimet"] = yhteenveto.get("muut_toimet", "")
    st.session_state["lisahuomiot"] = yhteenveto.get("lisahuomiot", "")
    st.session_state["tarkastaja_kuittaus"] = kuittaukset.get("tarkastaja_kuittaus", "")
    st.session_state["vastaanottaja_kuittaus"] = kuittaukset.get("vastaanottaja_kuittaus", "")

    sections = list(data.get("tarkastusosiot", {}).keys())
    transformer_sections = [sec for sec in sections if sec.startswith("Muuntaja ")]
    st.session_state["transformer_count"] = len(transformer_sections) if transformer_sections else 1
    st.session_state["selected_sections"] = sections if sections else [sec for sec in DEFAULT_ENABLED if sec in get_available_sections()]

    for section_name, rows in data.get("tarkastusosiot", {}).items():
        for idx, row in enumerate(rows):
            st.session_state[make_key(section_name, idx, "status")] = row.get("tila", "Ei tarkastettu")
            st.session_state[make_key(section_name, idx, "risk")] = row.get("riskitaso", "Matala")
            st.session_state[make_key(section_name, idx, "comment")] = row.get("kommentti", "")
            st.session_state[make_key(section_name, idx, "action")] = row.get("toimenpide", "")
            st.session_state[make_key(section_name, idx, "image_data")] = row.get("kuvat", [])


init_state()

if st.session_state.pending_loaded_json is not None:
    load_json_to_state(st.session_state.pending_loaded_json)
    st.session_state.pending_loaded_json = None
    st.session_state.json_loaded_message = "Pöytäkirja ladattu lomakkeelle."
    st.rerun()

logo_bytes = load_logo_bytes()
if logo_bytes:
    st.image(logo_bytes, width=180)

st.title("Tarkastuspöytäkirja")
st.markdown(
    "Selainpohjainen Streamlit-sovellus sähkölaitteistojen säännöllisiin tarkastuskäynteihin. "
    "Tarkastusosiot valitaan asetuksista kohdekohtaisesti ja havaintoon voi liittää useita valokuvia."
)

with st.sidebar:
    st.header("Asetukset")

    st.divider()
    st.subheader("Tallenna / lataa")

    if st.session_state.json_loaded_message:
        st.success(st.session_state.json_loaded_message)
        st.session_state.json_loaded_message = ""

    uploaded_json = st.file_uploader("Lataa aiempi JSON-pöytäkirja", type=["json"])
    if uploaded_json is not None:
        try:
            loaded_data = json.load(uploaded_json)
            st.session_state.pending_loaded_json = loaded_data
            st.rerun()
        except Exception as exc:
            st.error(f"JSON-lataus epäonnistui: {exc}")

    st.number_input(
        "Muuntajien määrä",
        min_value=0,
        max_value=6,
        step=1,
        key="transformer_count",
        on_change=sync_selected_sections,
        help="Voit lisätä kohteelle useita muuntajaosioita, esim. Muuntaja 1 ja Muuntaja 2.",
    )

    available_sections = get_available_sections()
    default_selected = [sec for sec in st.session_state.get("selected_sections", DEFAULT_ENABLED) if sec in available_sections]
    st.multiselect(
        "Valitse tarkastusosiot",
        options=available_sections,
        key="selected_sections",
        default=default_selected,
    )

st.header("1. Kohdetiedot")
col1, col2, col3 = st.columns(3)

with col1:
    st.text_input("Kohteen nimi", key="kohteen_nimi")
    st.text_input("Osoite", key="osoite")
    st.date_input("Tarkastuspäivä", key="tarkastuspaiva", value=date.today(), format="DD.MM.YYYY")

with col2:
    st.text_input("Tarkastaja", key="tarkastaja")
    st.text_input("Käytönjohtaja", key="kaytonjohtaja")
    st.text_input("Kiinteistön yhteyshenkilö", key="yhteyshenkilo")

with col3:
    st.selectbox(
        "Suurin nimellisjännite",
        ["Pienoisjännite", "230 V", "400 V", "6 kV", "10 kV", "20 kV"],
        key="suurin_nimellisjannite",
    )
    st.text_input("Kohteen tyyppi", key="kohteen_tyyppi", placeholder="Esim. muuntamo / teollisuuskohde / pumppaamo")
    st.selectbox(
        "Tarkastuksen tyyppi",
        ["Muuntamotarkastus", "ATEX", "Muuntamotarkastus +ATEX"],
        key="tarkastuksen_tyyppi",
        on_change=sync_sections_from_tarkastustyyppi,
    )

st.text_input("Sää / olosuhteet", key="olosuhteet", placeholder="Esim. kuiva, pakkasta -8 C")

st.header("2. Tarkastusosiot")
selected_sections = st.session_state.get("selected_sections", [])
if st.session_state.get("tarkastuksen_tyyppi") in ["ATEX", "Muuntamotarkastus +ATEX"] and "ATEX-tila" not in selected_sections:
    selected_sections = selected_sections + ["ATEX-tila"]
    st.session_state.selected_sections = selected_sections

if not selected_sections:
    st.warning("Valitse vähintään yksi tarkastusosio asetuksista.")
else:
    for section in selected_sections:
        base_section = get_base_section_name(section)
        with st.expander(section, expanded=False):
            render_check_section(section, SECTION_DEFINITIONS[base_section])

st.header("3. Yhteenveto ja toimenpiteet")
col4, col5 = st.columns(2)

with col4:
    st.selectbox(
        "Kokonaisarvio",
        [
            "Laitteisto kunnossa",
            "Laitteistossa huomioita",
            "Korjaavia toimenpiteitä vaaditaan",
            "Välitön turvallisuusriski",
        ],
        key="kokonaisarvio",
    )
    st.text_area(
        "Välittömät toimenpiteet",
        key="valittomat_toimet",
        height=110,
        placeholder="Kirjaa heti tehtävät toimet, käytön rajoitukset tai erottamistarve.",
    )
    st.text_area(
        "Muut toimenpiteet",
        key="muut_toimet",
        height=110,
        placeholder="Kirjaa suunnitellut korjaukset, lisäselvitykset ja vastuuhenkilöt.",
    )

with col5:
    st.date_input("Seuraava tarkastus", key="seuraava_tarkastus", value=date.today(), format="DD.MM.YYYY")
    st.text_area(
        "Lisähuomiot",
        key="lisahuomiot",
        height=190,
        placeholder="Vapaa yhteenveto, rajaukset, suositukset ja lisätiedot.",
    )

st.header("4. Kuittaukset")
col6, col7 = st.columns(2)

with col6:
    st.text_input("Tarkastajan kuittaus", key="tarkastaja_kuittaus", placeholder="Esim. Nimi / päiväys")

with col7:
    st.text_input("Vastaanottajan kuittaus", key="vastaanottaja_kuittaus", placeholder="Esim. Nimi / päiväys")

st.divider()

report_data = gather_form_data()
json_bytes = json.dumps(report_data, ensure_ascii=False, indent=2).encode("utf-8")
html_report = build_report(report_data)
pdf_bytes = build_pdf(report_data)
word_bytes = build_word(report_data)

col8, col9, col10, col11, col12 = st.columns(5)

with col8:
    st.download_button(
        "Lataa JSON",
        data=json_bytes,
        file_name=f"tarkastuspoytakirja_{date.today().isoformat()}.json",
        mime="application/json",
        use_container_width=True,
    )

with col9:
    st.download_button(
        "Lataa HTML",
        data=html_report.encode("utf-8"),
        file_name=f"tarkastuspoytakirja_{date.today().isoformat()}.html",
        mime="text/html",
        use_container_width=True,
    )

with col10:
    st.download_button(
        "Lataa PDF",
        data=pdf_bytes,
        file_name=f"tarkastuspoytakirja_{date.today().isoformat()}.pdf",
        mime="application/pdf",
        use_container_width=True,
    )

with col11:
    st.download_button(
        "Lataa Word",
        data=word_bytes,
        file_name=f"tarkastuspoytakirja_{date.today().isoformat()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

with col12:
    st.download_button(
        "Lataa TXT",
        data=(
            f"Tarkastuspöytäkirja\n"
            f"Kohde: {report_data['kohde']['kohteen_nimi']}\n"
            f"Päivä: {report_data['kohde']['tarkastuspaiva']}\n"
            f"Tarkastaja: {report_data['kohde']['tarkastaja']}\n"
            f"Kokonaisarvio: {report_data['yhteenveto']['kokonaisarvio']}\n\n"
            f"Välittömät toimenpiteet:\n{report_data['yhteenveto']['valittomat_toimet']}\n\n"
            f"Muut toimenpiteet:\n{report_data['yhteenveto']['muut_toimet']}\n\n"
            f"Lisähuomiot:\n{report_data['yhteenveto']['lisahuomiot']}\n\n"
            f"{COPYRIGHT_TEXT}\n"
            f"{WEBSITE_URL}\n"
        ).encode("utf-8"),
        file_name=f"tarkastuspoytakirja_{date.today().isoformat()}.txt",
        mime="text/plain",
        use_container_width=True,
    )

st.subheader("Raporttiesikatselu")
st.components.v1.html(html_report, height=900, scrolling=True)

with st.expander("Asennus- ja käyttöohje"):
    st.markdown(
        f"""
Tämä tarkastuspöytäkirja toimii selaimessa puhelimella, tabletilla ja tietokoneella.

**Tarkastuksen tekeminen puhelimella**

1. Avaa tarkastuspöytäkirja puhelimen selaimessa.
2. Täytä kohdetiedot.
3. Valitse tarvittavat tarkastusosiot kohdasta **Asetukset**.
4. Kirjaa havainnot ja lisää tarvittaessa valokuvia havaintokohtiin.
5. Täytä lopuksi yhteenveto ja kuittaukset.

**JSON-tiedoston tallentaminen työn aikana**

JSON on ohjelman tallennustiedosto. Sen avulla tarkastusta voi jatkaa myöhemmin toisella laitteella.

1. Paina lopuksi tai työn aikana **Lataa JSON**.
2. Puhelin tallentaa tiedoston latauksiin.
3. Tiedoston voi lähettää itselle esimerkiksi sähköpostilla tai siirtää pilvipalveluun.

**JSON-tiedoston avaaminen myöhemmin tietokoneella**

1. Avaa tarkastuspöytäkirja tietokoneella selaimeen.
2. Avaa vasemman reunan **Asetukset**.
3. Kohdassa **Tallenna / lataa** valitse aiemmin tallennettu JSON-tiedosto.
4. Lomake täyttyy automaattisesti aiemmin tallennetuilla tiedoilla.
5. Viimeistele tarkastus ja lataa lopullinen raportti PDF-, Word-, HTML- tai TXT-muodossa.

**Suositeltu toimintatapa kenttätyössä**

- Tee tarkastus puhelimella paikan päällä.
- Tallenna JSON työn lopuksi tai jo välivaiheessa.
- Avaa sama JSON myöhemmin tietokoneella, jos haluat viimeistellä raportin suuremmalla näytöllä.
- Lataa lopullinen raportti vasta viimeistelyn jälkeen.

Lisätiedot: {WEBSITE_URL}
"""
    )
